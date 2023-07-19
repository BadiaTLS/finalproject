from datetime import time
import openpyxl
from django.http import HttpResponse
import os
from .models import table_session, table_time, table_live_booking
from django.db import transaction


def fetch_all_session_objects():
    return table_session.objects.all()

def fetch_time_objects(session_id):
    session = table_session.objects.get(id=session_id)
    return session.table_time_set.all()

def get_session_by_id(session_id):
    return table_session.objects.get(pk=session_id)

def get_time_by_id(time_id):
    return table_time.objects.get(pk=time_id)

def update_session_limit(form_session):
    form_session.save()

def update_session_available_seat(time_id, limit_after):
    time_object =  table_time.objects.get(pk=time_id)
    limit_before = time_object.seat_limit
    available_before = time_object.available_seat
    difference = abs(limit_before - limit_after)
    print(limit_before, limit_after, difference, available_before)
    print(time_object, type(time_object))

    if available_before < time_object.available_seat:
        return False

    if limit_before > limit_after: 
        time_object.available_seat = available_before - difference
        print(f"Lebih besar {limit_before, limit_after}")
    elif limit_before < limit_after: 
        value = (available_before + difference)
        time_object.available_seat = value
        print(f"Lebih kecil = {available_before} + {difference} = {time_object.available_seat}")
    else: 
        time_object.available_seat = time_object.seat_limit
    
    time_object.seat_limit = limit_after
    time_object.save()
    return True

def delete_session_object(session_id):
    session_id.delete()

def delete_time_object(time_id):
    time_id.delete()

def export_data_to_excel(file_path):
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    headers = ["Date", "Name", "Time", "Available", "Limit", "Menu"]
    for col_num, header in enumerate(headers, 1):
        worksheet.cell(row=1, column=col_num, value=header)

    sessions = fetch_all_session_objects()
    for session in sessions:
        times = fetch_time_objects(session.id)

        for time_obj in times:
            date = session.date
            name = session.get_name_display()
            menu = session.menu

            time_value = time_obj.time
            available = time_obj.available_seat
            limit = time_obj.seat_limit

            row_data = [date, name, time_value, available, limit, menu]
            row_num = worksheet.max_row + 1
            for col_num, value in enumerate(row_data, 1):
                worksheet.cell(row=row_num, column=col_num, value=value)

    workbook.save(file_path)

def save_session_and_times(form_session, form_time):
    session = form_session.save()

    if session.name == "Breakfast":
        times = [
            time(7, 0),
            time(7, 30),
            time(8, 0),
            time(8, 30),
        ]
    elif session.name == "Lunch":
        times = [
            time(11, 0),
            time(11, 30),
            time(12, 0),
            time(12, 30),
            time(13, 0),
            time(13, 30),
        ]
    elif session.name == "Dinner":
        times = [
            time(17, 0),
            time(17, 30),
            time(18, 0),
            time(18, 30),
            time(19, 0),
            time(19, 30),
        ]
    else:
        times = []

    for time_value in times:
        table_time.objects.create(
            time=time_value,
            session_id=session,
            seat_limit=form_time.cleaned_data["seat_limit"],
            available_seat = form_time.cleaned_data["seat_limit"],
        )

def download_file_response(file_path):
    with open(file_path, 'rb') as file:
        file_data = file.read()

    response = HttpResponse(file_data, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="order_record.xlsx"'

    os.remove(file_path)
    return response

# Upload Session Menu START
def read_excel_file(excel_file):
    try:
        workbook = openpyxl.load_workbook(excel_file)
        worksheet = workbook.active
    except Exception as e:
        raise Exception(f'Error reading Excel file: {e}')
    
    return worksheet

def create_times(session):
    if session.name == "Breakfast":
        times = [
            time(7, 0),
            time(7, 30),
            time(8, 0),
            time(8, 30),
        ]
    elif session.name == "Lunch":
        times = [
            time(11, 0),
            time(11, 30),
            time(12, 0),
            time(12, 30),
            time(13, 0),
            time(13, 30),
        ]
    elif session.name == "Dinner":
        times = [
            time(17, 0),
            time(17, 30),
            time(18, 0),
            time(18, 30),
            time(19, 0),
            time(19, 30),
        ]
    else:
        times = []
    
    return times

def update_or_create_table_time(session, time, seat_limit):
    return table_time.objects.update_or_create(
        session_id=session,
        time=time,
        defaults={
            'seat_limit': seat_limit,
            'available_seat': seat_limit,
        }
    )

# Upload Session Menu END

# DOWNLOAD REPORT FUNTIONS START #
import os
from django.http import HttpResponse
from datetime import datetime
from docx import Document

from io import BytesIO

def download_report_doc(start_date, end_date, filename):
    # Generate the Word document in memory
    document = Document()
    document.add_heading('Order Report', 0)

    headers = ["Date", "Name", "Time Range", "Available", "Limit", "Menu"]
    table = document.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    sessions = fetch_all_session_objects()
    for session in sessions:
        if start_date <= session.date <= end_date:
            times = fetch_time_objects(session.id)

            # Calculate start and end times for this session
            start_time = min(time_obj.time for time_obj in times)
            end_time = max(time_obj.time for time_obj in times)
            time_range = f"{start_time} - {end_time}"

            # Get other data for this session
            date = session.date
            name = session.get_name_display()
            menu = session.menu
            available = sum(time_obj.available_seat if time_obj.available_seat is not None else time_obj.seat_limit for time_obj in times)
            limit = sum(time_obj.seat_limit for time_obj in times)

            # Add row to table
            row_data = [date, name, time_range, available, limit, menu]
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)

    # Save the Word document to a BytesIO object
    file_data = BytesIO()
    document.save(file_data)

    # Set the response headers and content
    response = HttpResponse(file_data.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response

def validate_dates(start_date: str, end_date: str, date_format: str = '%Y-%m-%d') -> bool:
    """
    Validate if start_date and end_date are valid dates in the specified format.

    :param start_date: The start date string to validate.
    :param end_date: The end date string to validate.
    :param date_format: The date format to use for validation (default: '%Y-%m-%d').
    :return: True if both start_date and end_date are valid dates in the specified format, False otherwise.
    """
    try:
        parsed_start_date = datetime.strptime(start_date, date_format).date()
        parsed_end_date = datetime.strptime(end_date, date_format).date()
        return True, parsed_start_date, parsed_end_date
    except ValueError:
        return False
    
# DOWNLOAD REPORT FUNTIONS END #

# Upload Live Booking START
def process_upload(request):
    excel_file = request.FILES['excel_file']

    if excel_file.name.endswith('.xlsx'):
        try:
            workbook = openpyxl.load_workbook(excel_file)
            worksheet = workbook.active
        except Exception as e:
            return {'error': f'Error reading Excel file: {e}'}

        # validate required columns
        required_columns = ['Arrival Time', 'Served Time', 'Depart Time', 'Booking ID']
        headers = [cell.value for cell in worksheet[1]]
        missing_columns = [col for col in required_columns if col not in headers]
        if missing_columns:
            return {'error': f'Missing required columns: {", ".join(missing_columns)}'}

        with transaction.atomic():
            for row in worksheet.iter_rows(min_row=2, values_only=True):
                arrival_time = row[0]
                served_time = row[1]
                depart_time = row[2]
                bookings_id = row[3]
                if table_booking_dininghall.objects.filter(id=bookings_id).exists():
                    bookings, created = table_live_booking.objects.update_or_create(
                        arrival_time=arrival_time,
                        served_time=served_time,
                        depart_time=depart_time,
                        bookings_id_id=bookings_id
                    )

        return {'success': 'Live Booking imported successfully.'}
    else:
        return {'error': 'Invalid file format. Please upload an Excel file (.xlsx).'}


def get_upload_live_booking_file_context(request):
    return {'email': request.user.email}
# Upload Live Booking END

# Dashboard Context START
def get_dashboard_context(request):
    current_date = datetime.now()
    year = current_date.year
    month = current_date.month
    day = current_date.day

    recent_bookings = table_booking_dininghall.objects.order_by('-created_at')[:5]

    bar_data = get_bar_chart_data()

    line_data = get_line_chart_info()

    major_data = get_major_chart_info(date=current_date)

    todays_bookings_count = get_bookings_count(date=current_date)
    today_remaining_seats = get_remaining_seats(date=current_date)
    todays_bookings_count, today_remaining_seats = get_total_seat_info(date=current_date)
    
    total_remaining_seats, total_bookings_count = get_total_seat_info()

    todays_most_popular_session_name, todays_most_popular_session_count  = get_most_popular_session_info(date=current_date)

    avg_queue_time = get_average_queue_time(current_date)
    avg_dining_time = get_average_dining_time(current_date)
    average_stay_time = get_average_stay_time(current_date)
    average_queue_time_data = get_average_n_queue_time_chart_info(date=current_date)
    average_queue_time_per_session = get_average_session_queue_time_chart_info(date=current_date)

    x_lr, y_lr, x_lr_p, y_lr_p, mad, mse, mape = get_lr_data(date=current_date)

    context = {
        'antrian_n_hari' : average_queue_time_data,
        'antrian_per_sesi' : average_queue_time_per_session, 
        'bar_data' : bar_data,
        'line_data' : line_data,
        'major_data': major_data,
        'todays_booking_count' : todays_bookings_count,
        'todays_remaining_seats' : today_remaining_seats,
        'total_remaining_seats' : total_remaining_seats, 
        'total_bookings_count' : total_bookings_count,
        'todays_most_popular_session_name': todays_most_popular_session_name,
        'todays_most_popular_session_count' : todays_most_popular_session_count,
        'recent_bookings': recent_bookings,
        'avg_queue_time': avg_queue_time,
        'avg_dining_time': avg_dining_time, 
        'average_stay_time': average_stay_time,
        'x_lr': x_lr,
        'y_lr': y_lr,
        'x_lr_p': x_lr_p,
        'y_lr_p': y_lr_p,
        'mad' : mad,
        'mse': mse,
        'mape': mape,
        'email' : request.user.email,
    }
    return context

# Dashboard Context END

### GET CHART DATA START ###
def get_average_n_queue_time_chart_info(date, n = 7):
    labels = []
    queues_data = []
    # Get today's date

    # Calculate the start date as n days before the given date
    start_date = date - timedelta(days=n)

    # Create a list of dates for the last n days
    dates = [start_date + timedelta(days=i) for i in range(n+1)]


    for date in dates:
        labels.append(date.strftime('%Y-%m-%d'))
        queues_data.append(get_average_queue_time_in_seconds(target_date=date))

    data = {
        'labels': labels,
        'datasets': [
            {
                'label': 'Queue Time in Seconds',
                'data': queues_data,
                'backgroundColor': 'rgba(255, 255, 0, .1)',
                'fill': True
            },
        ]
    }
    data_json = json.dumps(data)
    return data_json

### GET LR Data ###
def get_lr_data(date, n=30):
    original_dates = []
    original_values = []

    # Calculate the start date as n days before the given date
    start_date = date - timedelta(days=n)

    # Create a list of dates for the last n days (excluding Saturdays and Sundays)
    for i in range(n + 1):
        current_date = start_date + timedelta(days=i)
        if current_date.weekday() < 5:  # Monday is 0, Friday is 4 (0 to 4 are weekdays)
            original_dates.append(current_date.strftime('%Y-%m-%d'))
            original_values.append(get_average_queue_time_in_seconds(target_date=current_date))

    x = ['2023-06-19', '2023-06-20', '2023-06-21', '2023-06-22', '2023-06-23', '2023-06-26', '2023-06-27', '2023-06-28', '2023-06-29', '2023-06-30', '2023-07-03', '2023-07-04', '2023-07-05', '2023-07-06', '2023-07-07', '2023-07-10', '2023-07-11', '2023-07-12', '2023-07-13', '2023-07-14', '2023-07-17', '2023-07-18', '2023-07-19'] 
    y = [43, 41, 46, 44, 44, 43, 47, 44, 43, 45, 45, 43, 46, 46, 43, 43, 45, 41, 45, 44, 44, 44, 45]
    window_size = 3
    num_steps = 7

    # Predict the queue time for the next 'predicted_n' days using Moving Average
    predicted_x_values, predicted_y_values = moving_average_forecast(original_dates, original_values, window_size=window_size, num_steps=num_steps)


    # MA Evaluation
    middle = len(original_values) // 2
    training = original_values[:middle]
    test = original_values[middle:]

    predictions = sma(training)
    mad, mse, mape = evaluate(predictions, test)
    print(f"MAD: {mad:.2f}, MSE: {mse:.2f}, MAPE: {mape:.2%}")
    mad = f"{mad:.2f}"
    mse = f"{mse:.2f}"
    mape = f"{mape:.2%}"

    middle = len(original_values) // 2
    # original_values[:middle] = [None] * middle


    # Convert the values to JSON for plotting in HTML
    xnumbersJson = json.dumps(original_dates)
    ynumbersJson = json.dumps(original_values)
    predicted_x_numbers_json = json.dumps(predicted_x_values)
    predicted_y_numbers_json = json.dumps(predicted_y_values)
    return xnumbersJson, ynumbersJson, predicted_x_numbers_json, predicted_y_numbers_json, mad, mse, mape


### Predict MA Data ###
def sma(values, n=5, m=3):
    """Calculate the simple moving average of the m most recent values in a list and predict the next n values."""
    recent_values = values[-m:]
    avg = sum(recent_values) / len(recent_values)
    return [avg] * n

def evaluate(predictions, actuals):
    """Evaluate the performance of a prediction method by calculating the MAD, MSE, and MAPE."""
    errors = [p - a for p, a in zip(predictions, actuals)]
    mad = sum(abs(e) for e in errors) / len(errors)
    mse = sum(e ** 2 for e in errors) / len(errors)
    mape = sum(abs(e / a) for e, a in zip(errors, actuals)) / len(errors)
    return mad, mse, mape

import pandas as pd
def moving_average_forecast(dates, values, window_size, num_steps):
    # Calculate the moving average for the given window size
    moving_average = sum(values[-window_size:]) / window_size

    # Predict the next few days using the last available moving average value
    training_dates = dates[-window_size:]
    training_values = [int(moving_average)] * window_size

    print("TRAININD DATES AND VALUE", training_dates, training_values)

    # Predict the next few days using the last available moving average value
    test_future_dates = pd.date_range(start=dates[-1], periods=num_steps+1)[1:]
    test_future_values = [int(moving_average)] * num_steps

    return [date.strftime('%Y-%m-%d') for date in test_future_dates], test_future_values



from datetime import timedelta
import json

def get_average_session_queue_time_chart_info(date, n=0, selected_session=None):
    labels = []
    if selected_session:
        sessions = [selected_session]
    else:
        sessions = [table_session.breakfast, table_session.lunch, table_session.dinner]

    session_hours = {
        table_session.breakfast: range(7, 9),
        table_session.lunch: range(11, 14),
        table_session.dinner: range(17, 20)
    }
    session_data = {session: [[] for _ in range(48)] for session in sessions}
    base_colors = ["rgba(0, 156, 255, .99)", "rgba(255, 99, 132, .99)", "rgba(75, 192, 192, .99)"]

    start_date = date - timedelta(days=n)
    dates = [start_date + timedelta(days=i) for i in range(n+1)]

    for date in dates:
        labels.append(date.strftime('%Y-%m-%d'))
        for session in sessions:
            queue_times = get_average_queue_time_in_seconds_for_session(target_date=date,session_hours=session_hours , session=session)
            for i, queue_time in enumerate(queue_times):
                hour = i // 2
                if hour not in session_hours[session]:
                    continue
                session_data[session][i].append(queue_time)

    datasets = []
    for session, base_color in zip(sessions, base_colors):
        color_step = int(255 / (len(session_hours[session]) * 2))
        for i in range(48):
            hour = i // 2
            if hour not in session_hours[session]:
                continue
            minute = i % 2 * 30
            color_index = list(session_hours[session]).index(hour) * 2 + minute // 30
            color_value = color_index * color_step
            r, g, b, a = [int(x) if x.isnumeric() else float(x) for x in base_color[5:-1].split(',')]
            alpha = max(0.2, a - color_value / 255)
            color = f"rgba({r}, {g}, {b}, {alpha})"
            datasets.append({
                'label': f"{session} - {hour}:{minute:02d}",
                'data': session_data[session][i],
                "backgroundColor": color,
                'fill': False
            })

    data = {
        'labels': labels,
        'datasets': datasets
    }
    data_json = json.dumps(data)
    return data_json


def get_major_chart_info(date, n=0):
    # Calculate the start date as n days before the given date
    start_date = date - timedelta(days=n)

    # Calculate the average queue time for each major
    results = table_live_booking.objects.filter(
        bookings_id__session_id__date__range=(start_date, date),
        served_time__isnull=False,
        arrival_time__isnull=False,
    ).annotate(
        queue_time=F('served_time') - F('arrival_time')
    ).values('bookings_id__user_id__major').annotate(avg_queue_time=Avg('queue_time'))

    # Format the results as a dictionary
    data = {result['bookings_id__user_id__major']: result['avg_queue_time'].total_seconds() for result in results}

    # You can then use this data to populate the 'data' field in your chart
    label = ["IBDA", "IEE", "CFP", "BMS", "SCCE", "ASD"]
    majors = ["ibda", "iee", "cfp", "bms", "scce", "asd"]
    chart_data = {
        'labels': label,
        'datasets': [{
            'backgroundColor': [
                "rgba(255, 99, 132, 0.7)",
                "rgba(54, 162, 235, 0.7)",
                "rgba(255, 206, 86, 0.7)",
                "rgba(75, 192, 192, 0.7)",
                "rgba(153, 102, 255, 0.7)",
                "rgba(255, 159, 64, 0.7)"
            ],
            'data': [data.get(major, 0) for major in majors]
        }]
    }
    data_json = json.dumps(chart_data)
    return data_json

from .models import table_booking_dininghall
from django.db.models import Count, Sum
from datetime import date, timedelta
import json

def get_bar_chart_data():
    # Get the start and end dates for the week
    today = date.today()
    start_date = today - timedelta(days=today.weekday())
    end_date = start_date + timedelta(days=6)

    # Query the database for the bookings
    bookings = table_booking_dininghall.objects.filter(
        session_id__date__range=(start_date, end_date)
    ).values(
        'session_id__date', 'session_id__name'
    ).annotate(
        count=Count('id')
    ).order_by(
        'session_id__date', 'session_id__name'
    )

    # Initialize the data for the chart
    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    sessions = [table_session.breakfast, table_session.lunch, table_session.dinner]
    colors = ["rgba(0, 156, 255, .7)", "rgba(255, 99, 132, .7)", "rgba(75, 192, 192, .7)"]
    data = {
        "labels": days,
        "datasets": [
            {
                "label": session,
                "data": [0] * len(days),
                "backgroundColor": color
            } for session, color in zip(sessions, colors)
        ]
    }

    # Fill in the data for the chart
    for booking in bookings:
        day_index = booking['session_id__date'].weekday()
        session_index = sessions.index(booking['session_id__name'])
        data['datasets'][session_index]['data'][day_index] = booking['count']
    data_json = json.dumps(data)
    return data_json


def get_line_chart_info():
    labels = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    reserved_data = []
    unreserved_data = []
    # Get today's date
    today = datetime.today()

    # Calculate the start date of the week (Monday)
    start_of_week = today - timedelta(days=today.weekday())

    # Create a list of dates for the week
    week_dates = [start_of_week + timedelta(days=i) for i in range(7)]

    for label, date in zip(labels, week_dates):
        total_remaining_seats, total_bookings_count = get_total_seat_info(date=date)
        reserved_data.append(total_bookings_count)
        unreserved_data.append(total_remaining_seats)

    data = {
        'labels': labels,
        'datasets': [
            {
                'label': 'Reserved',
                'data': reserved_data,
                'backgroundColor': 'rgba(255, 255, 0, .1)',
                'fill': True
            },
            {
                'label': 'Unreserved',
                'data': unreserved_data,
                'backgroundColor': 'rgba(255, 99, 132, .7)',
                'fill': True
            }
        ]
    }
    data_json = json.dumps(data)
    return data_json

## 
from django.utils import timezone
from .models import table_booking_dininghall

def get_bookings_count(date):
    todays_bookings_count = table_booking_dininghall.objects.filter(session_id__date=date).count()
    return todays_bookings_count

from django.utils import timezone
from .models import table_booking_dininghall, table_time

def get_remaining_seats(date):
    seat_limit = table_time.objects.filter(session_id__date=date).aggregate(Sum('seat_limit'))['seat_limit__sum']
    bookings_count = table_booking_dininghall.objects.filter(session_id__date=date).count()
    if seat_limit is not None:
        remaining_seats = seat_limit - bookings_count
    else:
        remaining_seats = 0 # or another default value that makes sense for your program
    return remaining_seats

def get_total_seat_info11():
    total_seat_limit = table_time.objects.aggregate(Sum('seat_limit'))['seat_limit__sum']
    total_bookings_count = table_booking_dininghall.objects.count()
    total_remaining_seats = total_seat_limit - total_bookings_count
    return total_remaining_seats, total_bookings_count

def get_total_seat_info(date=None, year=None, month=None, day=None):
    table_time_qs = table_time.objects.all()
    table_booking_dininghall_qs = table_booking_dininghall.objects.all()
    
    if date:
        table_time_qs = table_time_qs.filter(session_id__date=date)
        table_booking_dininghall_qs = table_booking_dininghall_qs.filter(session_id__date=date)
    
    if year:
        table_time_qs = table_time_qs.filter(session_id__date__year=year)
        table_booking_dininghall_qs = table_booking_dininghall_qs.filter(session_id__date__year=year)
    
    if month:
        table_time_qs = table_time_qs.filter(session_id__date__month=month)
        table_booking_dininghall_qs = table_booking_dininghall_qs.filter(session_id__date__month=month)
    
    if day:
        table_time_qs = table_time_qs.filter(session_id__date__day=day)
        table_booking_dininghall_qs = table_booking_dininghall_qs.filter(session_id__date__day=day)
    

    total_seat_limit = table_time_qs.aggregate(Sum('seat_limit'))['seat_limit__sum']
    if total_seat_limit is None:
        total_seat_limit = 0
    total_bookings_count = table_booking_dininghall_qs.count()
    total_remaining_seats = total_seat_limit - total_bookings_count
    return total_remaining_seats, total_bookings_count

from django.db.models import Count
from django.utils import timezone
from .models import table_session

def get_most_popular_session_info(date):
    today = date
    most_popular_session = table_session.objects.filter(date=today).annotate(bookings_count=Count('table_booking_dininghall')).order_by('-bookings_count').first()
    if most_popular_session:
        session_name = most_popular_session.name
        bookings_count = most_popular_session.bookings_count
        return session_name, bookings_count
    else:
        return None, 0
    
from django.db.models import Avg, F, ExpressionWrapper
from django.db.models.functions import Cast
from datetime import timedelta

def get_average_queue_time(target_date):
    # Filter the table_live_booking queryset by the desired date, session, and served_time
    queue_times = table_live_booking.objects.filter(
        bookings_id__session_id__date=target_date,
        served_time__isnull=False
    )

    # Calculate the average queue time
    average_queue_time = queue_times.aggregate(
        avg_queue_time=Avg(F('served_time') - F('arrival_time'))
    )['avg_queue_time']

    # Format the average queue time as a string with only the minute and second components
    if average_queue_time is not None:
        total_seconds = int(average_queue_time.total_seconds())
        minutes, seconds = divmod(total_seconds, 60)
        average_queue_time_str = f"{minutes}m {seconds}s"
    else:
        average_queue_time_str = None

    return average_queue_time_str


def get_average_dining_time(target_date):
    # Filter the table_live_booking queryset by the desired date, session, and depart_time
    dining_times = table_live_booking.objects.filter(
        bookings_id__session_id__date=target_date,
        served_time__isnull=False
    )

    # Calculate the average dining time
    average_dining_time = dining_times.aggregate(
        avg_dining_time=Avg(F('depart_time') - F('served_time'))
    )['avg_dining_time']

    # Format the average dining time as a string with only the minute and second components
    if average_dining_time is not None:
        total_seconds = int(average_dining_time.total_seconds())
        minutes, seconds = divmod(total_seconds, 60)
        average_dining_time_str = f"{minutes}m {seconds}s"
    else:
        average_dining_time_str = None

    return average_dining_time_str

def get_average_stay_time(target_date: date):
    # Filter the table_live_booking queryset by the desired date, session, and depart_time
    stay_times = table_live_booking.objects.filter(
        bookings_id__session_id__date=target_date,
        depart_time__isnull=False
    )

    # Calculate the average stay time
    average_stay_time = stay_times.aggregate(
        avg_stay_time=Avg(F('depart_time') - F('arrival_time'))
    )['avg_stay_time']

    # Format the average stay time as a string with only the minute and second components
    if average_stay_time is not None:
        total_seconds = int(average_stay_time.total_seconds())
        minutes, seconds = divmod(total_seconds, 60)
        average_stay_time_str = f"{minutes}m {seconds}s"
    else:
        average_stay_time_str = None

    return average_stay_time_str

def get_average_dining_time_in_seconds(target_date):
    # Filter the table_live_booking queryset by the desired date, session, and depart_time
    dining_times = table_live_booking.objects.filter(
        bookings_id__session_id__date=target_date,
        served_time__isnull=False
    )

    # Calculate the average dining time
    average_dining_time = dining_times.aggregate(
        avg_dining_time=Avg(F('depart_time') - F('served_time'))
    )['avg_dining_time']

    # Calculate the total number of seconds for the average dining time
    if average_dining_time is not None:
        total_seconds = int(average_dining_time.total_seconds())
    else:
        total_seconds = None

    return total_seconds

def get_average_dining_time_in_minutes(target_date):
    # Filter the table_live_booking queryset by the desired date, session, and depart_time
    dining_times = table_live_booking.objects.filter(
        bookings_id__session_id__date=target_date,
        served_time__isnull=False
    )

    # Calculate the average dining time
    average_dining_time = dining_times.aggregate(
        avg_dining_time=Avg(F('depart_time') - F('served_time'))
    )['avg_dining_time']

    # Calculate the total number of minutes for the average dining time
    if average_dining_time is not None:
        total_seconds = int(average_dining_time.total_seconds())
        total_minutes = total_seconds // 60
    else:
        total_minutes = None

    return total_minutes

def get_average_queue_time_in_seconds(target_date, session=False):
    # Filter the table_live_booking queryset by the desired date, session, and served_time
    queue_times = table_live_booking.objects.filter(
    bookings_id__session_id__date=target_date,
    served_time__isnull=False,
    )
    if session:
        queue_times = queue_times.filter(bookings_id__session_id__name=session)

    # Calculate the average queue time
    average_queue_time = queue_times.aggregate(
        avg_queue_time=Avg(F('served_time') - F('arrival_time'))
    )['avg_queue_time']

    # Calculate the total number of minutes for the average queue time
    if average_queue_time is not None:
        total_seconds = int(average_queue_time.total_seconds())
    else:
        total_seconds = 0

    return total_seconds

def get_average_queue_time_in_seconds_for_session(target_date, session_hours, session=None):
    average_queue_times = []
    for hour in range(24):
        for minute in [0, 30]:
            if session and hour not in session_hours[session]:
                average_queue_times.append(0)
                continue

            # Filter the table_live_booking queryset by the desired date, session, served_time, hour, and minute
            queue_times = table_live_booking.objects.filter(
                bookings_id__session_id__date=target_date,
                served_time__isnull=False,
                arrival_time__hour=hour,
                arrival_time__minute__gte=minute,
                arrival_time__minute__lt=minute+30
            )
            if session:
                queue_times = queue_times.filter(bookings_id__session_id__name=session)

            # Calculate the average queue time
            average_queue_time = queue_times.aggregate(
                avg_queue_time=Avg(F('served_time') - F('arrival_time'))
            )['avg_queue_time']

            # Calculate the total number of minutes for the average queue time
            if average_queue_time is not None:
                total_seconds = int(average_queue_time.total_seconds())
            else:
                total_seconds = 0

            average_queue_times.append(total_seconds)

    return average_queue_times


### GET CHART DATA END ###